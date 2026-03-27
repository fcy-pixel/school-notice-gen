"""
Word 範本填充服務
使用 python-docx 將佔位符替換為通告內容
"""

import re
import os
from copy import deepcopy
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")


def ensure_templates_dir():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)


def list_templates() -> list[dict]:
    ensure_templates_dir()
    templates = []
    for filename in os.listdir(TEMPLATES_DIR):
        if filename.endswith(".docx"):
            templates.append({"id": filename, "name": filename.replace(".docx", "")})
    return templates


def get_template_path(template_id: str) -> str:
    path = os.path.join(TEMPLATES_DIR, template_id)
    if not os.path.exists(path):
        raise FileNotFoundError(f"範本不存在：{template_id}")
    return path


def extract_placeholders(template_id: str) -> list[str]:
    """提取範本中所有 {{佔位符}}"""
    path = get_template_path(template_id)
    doc = Document(path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text += para.text
    placeholders = re.findall(r"\{\{(.+?)\}\}", text)
    return list(set(placeholders))


def _para_full_text(para) -> str:
    """取得段落的完整文字（包含跨 run 的情況）。"""
    return "".join(t.text or "" for t in para._p.iter(qn("w:t")))


def _set_para_text(para, new_text: str):
    """
    將段落所有 run 清空，只保留第一個 run 並寫入 new_text。
    若無 run，則直接新增一個。
    """
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
    else:
        # 段落本身沒有 w:r 元素，手動新增
        r_elem = etree.SubElement(para._p, qn("w:r"))
        t_elem = etree.SubElement(r_elem, qn("w:t"))
        t_elem.text = new_text
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _replace_in_paragraph(para, replacements: dict[str, str]):
    """單行替換：value 不含換行。"""
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = _para_full_text(para)
        if placeholder not in full_text:
            continue
        new_text = full_text.replace(placeholder, value)
        # 先嘗試逐 run 替換（保留各 run 格式）
        replaced = False
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                replaced = True
        # 若佔位符跨多 run，回退到整段重寫
        if not replaced:
            _set_para_text(para, new_text)


def _expand_multiline_in_doc(doc, key: str, value: str):
    """
    將包含換行的佔位符（{{key}}）展開成多個 Word 段落。
    按 \\n\\n 分段；若無雙換行則按 \\n 分行（每行一段）。
    """
    placeholder = f"{{{{{key}}}}}"
    # 找目標段落（只搜尋 doc.paragraphs）
    target_para = None
    for para in doc.paragraphs:
        if placeholder in _para_full_text(para):
            target_para = para
            break

    if target_para is None:
        return

    # 決定分段方式
    if "\n\n" in value:
        parts = [p for p in value.split("\n\n") if p.strip()]
    else:
        parts = [p for p in value.split("\n") if p.strip()]

    if not parts:
        _set_para_text(target_para, _para_full_text(target_para).replace(placeholder, ""))
        return

    # 把原始段落文字中的佔位符替換成第一段內容
    original_text = _para_full_text(target_para)
    first_text = original_text.replace(placeholder, parts[0])
    _set_para_text(target_para, first_text)

    # 在原段落後依序插入其餘段落（複製相同的段落格式）
    parent = target_para._p.getparent()
    insert_idx = list(parent).index(target_para._p)
    for i, part_text in enumerate(parts[1:]):
        new_p = deepcopy(target_para._p)
        # 清除新段落的所有 run
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        # 新增一個 run 放入文字
        r_elem = etree.SubElement(new_p, qn("w:r"))
        t_elem = etree.SubElement(r_elem, qn("w:t"))
        t_elem.text = part_text
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        parent.insert(insert_idx + 1 + i, new_p)


def fill_template(template_id: str, replacements: dict[str, str]) -> BytesIO:
    """
    將佔位符替換後，回傳填好的 .docx 位元組流
    """
    path = get_template_path(template_id)
    doc = Document(path)

    # 第一階段：含換行的值展開成多個段落（例如通告正文）
    multiline = {k: v for k, v in replacements.items() if "\n" in str(v)}
    for key, value in multiline.items():
        _expand_multiline_in_doc(doc, key, value)

    # 第二階段：單行替換（所有段落、表格、頁首頁尾）
    singleline = {k: v for k, v in replacements.items() if "\n" not in str(v)}

    for para in doc.paragraphs:
        _replace_in_paragraph(para, singleline)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, singleline)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, singleline)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, singleline)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def create_default_template() -> str:
    """
    建立一個預設的測試範本（如果 templates/ 目錄為空）
    """
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 標題
    title = doc.add_heading("{{通告標題}}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # 收件人
    recipient_para = doc.add_paragraph()
    recipient_para.add_run("{{收件人}}")

    doc.add_paragraph()  # 空行

    # 正文
    doc.add_paragraph("{{正文}}")

    doc.add_paragraph()  # 空行

    # 落款
    closing_para = doc.add_paragraph()
    closing_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    closing_para.add_run("{{落款}}")

    path = os.path.join(TEMPLATES_DIR, "預設範本.docx")
    ensure_templates_dir()
    doc.save(path)
    return "預設範本.docx"
